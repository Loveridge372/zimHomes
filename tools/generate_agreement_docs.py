from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "docs" / "agreements"
OUTPUT_DIR = ROOT / "generated" / "agreements"


def add_markdown_line(document: Document, line: str) -> None:
    if not line.strip():
        document.add_paragraph()
        return

    if line.startswith("# "):
        paragraph = document.add_heading(line[2:].strip(), level=0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    if line.startswith("## "):
        document.add_heading(line[3:].strip(), level=1)
        return

    if line.startswith("### "):
        document.add_heading(line[4:].strip(), level=2)
        return

    if line.startswith("- "):
        document.add_paragraph(line[2:].strip(), style="List Bullet")
        return

    stripped = line.strip()
    if len(stripped) > 3 and stripped[0].isdigit() and ". " in stripped[:4]:
        document.add_paragraph(stripped.split(". ", 1)[1], style="List Number")
        return

    paragraph = document.add_paragraph()
    paragraph.add_run(stripped.replace("**", ""))


def build_docx(source_path: Path, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    for line in source_path.read_text(encoding="utf-8").splitlines():
        add_markdown_line(document, line)

    document.core_properties.author = "ZimHomes"
    document.core_properties.title = source_path.stem.replace("-", " ").title()
    document.save(output_path)


def markdown_to_pdf_story(source_path: Path) -> list:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="AgreementTitle",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            spaceAfter=16,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AgreementHeading",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#17221d"),
            spaceBefore=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AgreementBody",
            parent=styles["BodyText"],
            alignment=TA_LEFT,
            fontName="Helvetica",
            fontSize=9.2,
            leading=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AgreementList",
            parent=styles["BodyText"],
            leftIndent=14,
            firstLineIndent=-8,
            fontName="Helvetica",
            fontSize=9.2,
            leading=12,
            spaceAfter=4,
        )
    )

    story = []
    for line in source_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 4))
            continue
        if stripped.startswith("# "):
            story.append(Paragraph(stripped[2:].replace("**", ""), styles["AgreementTitle"]))
            continue
        if stripped.startswith("## "):
            story.append(Paragraph(stripped[3:].replace("**", ""), styles["AgreementHeading"]))
            continue
        if stripped.startswith("### "):
            story.append(Paragraph(stripped[4:].replace("**", ""), styles["AgreementHeading"]))
            continue
        if stripped.startswith("- "):
            story.append(Paragraph(f"- {stripped[2:].replace('**', '')}", styles["AgreementList"]))
            continue
        if len(stripped) > 3 and stripped[0].isdigit() and ". " in stripped[:4]:
            story.append(Paragraph(stripped.replace("**", ""), styles["AgreementList"]))
            continue
        story.append(Paragraph(stripped.replace("**", ""), styles["AgreementBody"]))
    return story


def build_pdf(source_path: Path, output_path: Path) -> None:
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title=source_path.stem.replace("-", " ").title(),
        author="ZimHomes",
    )
    document.build(markdown_to_pdf_story(source_path))


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for source_path in [
        SOURCE_DIR / "residential-lease-agreement.md",
        SOURCE_DIR / "property-management-agreement.md",
    ]:
        docx_output_path = OUTPUT_DIR / f"{source_path.stem}.docx"
        pdf_output_path = OUTPUT_DIR / f"{source_path.stem}.pdf"
        build_docx(source_path, docx_output_path)
        build_pdf(source_path, pdf_output_path)
        print(docx_output_path)
        print(pdf_output_path)


if __name__ == "__main__":
    main()
